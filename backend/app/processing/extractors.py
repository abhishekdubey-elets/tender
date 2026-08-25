"""Text extractors, one per document class.

Each extractor returns an :class:`ExtractionResult` recording the *method* used
and, where meaningful, a *confidence*. Heavy backends (pypdf, OCR) are imported
lazily and are injectable via :class:`ExtractionContext`, so the pipeline is
fully testable without PDF/OCR system binaries.

OCR is attempted only for scanned PDFs and only when an engine is supplied;
otherwise the document is returned with ``needs_ocr=True`` (never silently
dropped).
"""
from __future__ import annotations

import io
import json
from collections.abc import Callable
from dataclasses import dataclass

from app.processing.errors import ExtractionFailure, UnsupportedDocumentError
from app.processing.types import (
    DocClass,
    DocumentMetadata,
    ExtractionResult,
    Method,
    SourceFile,
)

# Backend signatures.
PdfTextBackend = Callable[[bytes], str]


@dataclass(slots=True)
class OcrResult:
    text: str
    confidence: float | None = None


OcrEngine = Callable[[bytes], OcrResult]

# Minimum non-whitespace characters for a PDF to be considered "has text layer".
PDF_TEXT_MIN_CHARS = 20


@dataclass(slots=True)
class ExtractionContext:
    pdf_text_backend: PdfTextBackend | None = None
    ocr_engine: OcrEngine | None = None


# --------------------------------------------------------------------------- #
# Default heavy backends (lazy)
# --------------------------------------------------------------------------- #
def _default_pdf_text(content: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _pdf_info(content: bytes) -> tuple[int | None, dict]:
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        info = reader.metadata or {}
        meta = {
            "title": getattr(info, "title", None),
            "author": getattr(info, "author", None),
        }
        return len(reader.pages), meta
    except Exception:
        return None, {}


# --------------------------------------------------------------------------- #
# Extractors
# --------------------------------------------------------------------------- #
def extract_html(source: SourceFile, ctx: ExtractionContext) -> ExtractionResult:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(source.content, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    title = soup.title.string.strip() if soup.title and soup.title.string else None
    text = "\n".join(line for line in soup.get_text("\n").splitlines() if line.strip())
    return ExtractionResult(
        text=text,
        method=Method.HTML_BS4,
        doc_class_final=DocClass.html,
        confidence=0.98,
        metadata=DocumentMetadata(title=title),
    )


def extract_json(source: SourceFile, ctx: ExtractionContext) -> ExtractionResult:
    try:
        data = json.loads(source.content.decode("utf-8"))
    except (ValueError, UnicodeDecodeError) as exc:
        raise ExtractionFailure(f"invalid JSON: {exc}") from exc
    text = json.dumps(data, ensure_ascii=False, indent=2, sort_keys=True)
    return ExtractionResult(
        text=text,
        method=Method.JSON,
        doc_class_final=DocClass.json,
        confidence=1.0,
        metadata=DocumentMetadata(extra={"json": data}),
    )


def extract_text(source: SourceFile, ctx: ExtractionContext) -> ExtractionResult:
    text = source.content.decode("utf-8", errors="replace")
    return ExtractionResult(
        text=text, method=Method.TEXT, doc_class_final=DocClass.text, confidence=1.0
    )


def extract_docx(source: SourceFile, ctx: ExtractionContext) -> ExtractionResult:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise UnsupportedDocumentError("python-docx not installed") from exc
    try:
        document = docx.Document(io.BytesIO(source.content))
    except Exception as exc:
        raise ExtractionFailure(f"could not read DOCX: {exc}") from exc

    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    props = document.core_properties
    meta = DocumentMetadata(
        title=props.title or None,
        author=props.author or None,
        created=props.created.isoformat() if props.created else None,
        modified=props.modified.isoformat() if props.modified else None,
    )
    return ExtractionResult(
        text="\n".join(parts),
        method=Method.DOCX,
        doc_class_final=DocClass.docx,
        confidence=1.0,
        metadata=meta,
    )


def extract_xlsx(source: SourceFile, ctx: ExtractionContext) -> ExtractionResult:
    from openpyxl import load_workbook

    try:
        wb = load_workbook(io.BytesIO(source.content), read_only=True, data_only=True)
    except Exception as exc:
        raise ExtractionFailure(f"could not read XLSX: {exc}") from exc
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"# {ws.title}")
        for row in ws.iter_rows(values_only=True):
            lines.append("\t".join("" if c is None else str(c) for c in row))
    return ExtractionResult(
        text="\n".join(lines),
        method=Method.XLSX,
        doc_class_final=DocClass.xlsx,
        confidence=1.0,
        metadata=DocumentMetadata(sheet_names=wb.sheetnames),
    )


def extract_legacy_office(source: SourceFile, ctx: ExtractionContext) -> ExtractionResult:
    # Legacy binary .doc/.xls need external tooling (antiword/LibreOffice); we do
    # not attempt an unreliable parse. Flagged, not discarded.
    raise UnsupportedDocumentError("legacy binary Office format (.doc/.xls) not supported")


def _pdf_text_confidence(text: str, page_count: int | None) -> float:
    chars = len(text.strip())
    pages = page_count or 1
    density = chars / pages
    if density < 100:
        return 0.6
    if density < 800:
        return 0.8
    return 0.95


def extract_pdf(source: SourceFile, ctx: ExtractionContext) -> ExtractionResult:
    backend = ctx.pdf_text_backend or _default_pdf_text
    try:
        text = backend(source.content)
    except Exception as exc:
        raise ExtractionFailure(f"could not read PDF: {exc}") from exc

    page_count, info = _pdf_info(source.content)
    meta = DocumentMetadata(
        title=info.get("title"), author=info.get("author"), page_count=page_count
    )

    if len(text.strip()) >= PDF_TEXT_MIN_CHARS:
        return ExtractionResult(
            text=text,
            method=Method.PDF_TEXT,
            doc_class_final=DocClass.pdf,
            confidence=_pdf_text_confidence(text, page_count),
            metadata=meta,
        )

    # Scanned PDF (image-only): OCR when we can, otherwise flag for later OCR.
    if ctx.ocr_engine is not None:
        ocr = ctx.ocr_engine(source.content)
        return ExtractionResult(
            text=ocr.text,
            method=Method.PDF_OCR,
            doc_class_final=DocClass.pdf_scanned,
            confidence=ocr.confidence,
            ocr_used=True,
            metadata=meta,
        )
    return ExtractionResult(
        text=None,
        method=Method.PDF_TEXT,
        doc_class_final=DocClass.pdf_scanned,
        confidence=None,
        needs_ocr=True,
        metadata=meta,
        warnings=["scanned PDF: OCR required but no engine configured"],
    )


_EXTRACTORS: dict[DocClass, Callable[[SourceFile, ExtractionContext], ExtractionResult]] = {
    DocClass.html: extract_html,
    DocClass.json: extract_json,
    DocClass.text: extract_text,
    DocClass.docx: extract_docx,
    DocClass.xlsx: extract_xlsx,
    DocClass.doc: extract_legacy_office,
    DocClass.xls: extract_legacy_office,
    DocClass.pdf: extract_pdf,
    DocClass.pdf_scanned: extract_pdf,
}


def get_extractor(doc_class: DocClass):
    extractor = _EXTRACTORS.get(doc_class)
    if extractor is None:
        raise UnsupportedDocumentError(f"no extractor for {doc_class.value}")
    return extractor
