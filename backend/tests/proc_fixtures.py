"""Representative document fixtures for processing tests."""
from __future__ import annotations

import io
import os

FIXTURES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "fixtures")


def load_fixture(name: str) -> bytes:
    with open(os.path.join(FIXTURES_DIR, name), "rb") as fh:
        return fh.read()


def make_docx() -> bytes:
    import docx

    document = docx.Document()
    document.core_properties.title = "Award Order"
    document.core_properties.author = "MoHUA"
    document.add_heading("Work Order", level=1)
    document.add_paragraph("Contract awarded to Beta Constructions Ltd.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Value"
    table.rows[0].cells[1].text = "7500000"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def make_xlsx() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Awards"
    ws.append(["awardee", "value_inr"])
    ws.append(["Gamma Systems Pvt Ltd", 3200000])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_blank_pdf() -> bytes:
    """A valid, single-page PDF with no text layer → represents a scanned PDF."""
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()
